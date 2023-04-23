from django.shortcuts import render,redirect
from AudioToText.models import * 
from .models import * 
from PyPDF2 import PdfReader
from django.contrib.auth import authenticate, login
from django.http import JsonResponse
import openai
import subprocess
from . import chatview
from . import cmd
# Create your views here.


def cut(response,seg):
    texte = response
    taille_morceau = int(seg)

    morceaux = []
    debut_morceau = 0
    fin_morceau = taille_morceau

    while debut_morceau < len(texte):
        morceaux.append(texte[debut_morceau:fin_morceau])
        debut_morceau += taille_morceau
        fin_morceau += taille_morceau
        
    return morceaux

# def pdfExtractText(file):
#     reader = PdfReader(file)
#     number_of_pages = len(reader.pages)
#     page = reader.pages[0]
#     text = page.extract_text()
#     return text 
def pdfExtractText(file):
    reader = PdfReader(file)
    number_of_pages = len(reader.pages)
    text = ""
    for i in range(number_of_pages):
        page = reader.pages[i]
        text += page.extract_text()
    return text


def docExtractText(file):
    import docx

    doc = docx.Document(file)
    text = []

    for paragraph in doc.paragraphs:
        text.append(paragraph.text)

    return '\n'.join(text)
# def docExtractText(file):
#     import docx

#     doc = docx.Document(file)
#     text = []

#     for page in doc.pages:
#         for paragraph in page.paragraphs:
#             text.append(paragraph.text)
    
#     print( '\n'.join(text))
#     return '\n'.join(text)


def newR(request):
    a,b=1,0
    try:
        user = authenticate(username='cime', password='cime')
        login(request, user)
    except:
        pass
    reportA = Report.objects.filter(user_id=request.user.id)
    if request.method=="POST":
        audio_file = request.FILES.get('file')
        d = datetime.datetime.now()
        report = Report.objects.create(filename=str(d.minute)+"_"+audio_file.name,user=request.user,file=audio_file)
        report.save()
        # print(request.POST.get('file'))
        return redirect("NewRep")
    context = {'a':a,'b':b,'report':reportA}
    return render(request,"report.html",context) 

# def NewRep(request):
#     a,b=1,0
#     report = Report.objects.filter(user_id=request.user.id).last
#     return render(request,"newRep.html",{'report':report,'a':a,'b':b})

def NewRep(request):
    a,b=1,0
    report = Report.objects.filter(user_id=request.user.id).last()
    # try:
    #     cmd.execBuild("/home/cime/Iris"+report.file.url)
    # except:
    #     pass
    return redirect("chat")



def fonction_report(report,valeur):
    ext = report.filename.split(".")
    tmp = ext[-1:]
    ex = str(tmp[0])
    desc = ["centre d'intérêt pour un journaliste","sujet de reportage"]
    media_file_path = "/home/cime/Iris"+report.file.url
    subprocess.call(['chmod', '777', media_file_path])
    if ex == "pdf" : 
        response = pdfExtractText(media_file_path).replace("\n"," ")
    else:
        response = docExtractText(media_file_path).replace("\n"," ")
    
    try : 
        if len(response) <=200:
            completion = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "Tu es un écrivains professionnel"},
                {"role": "user", "content": "resume moi en langage '"+valeur+"' dans un langage beau, correct et fluide le texte :"+response},
            ]
            )
            gpt = completion["choices"][0]["message"]["content"]
            
        else:
            morceaux = cut(response,200)
            gpt=""
            for i in morceaux:
                completion = openai.ChatCompletion.create(
                    model="gpt-3.5-turbo",
                    
                    messages=[
                        {"role": "system", "content": "Tu es un journaliste professionnel"},
                        {"role": "user", "content": "resume moi en langage '"+valeur+"' dans un langage beau, correct et fluide le texte :"+i},
                    ])
                gpt = gpt + " " + completion["choices"][0]["message"]["content"]
            
        Res = Resume.objects.create(Report=report,text=gpt)
        Res.save()
        
        resultat = Res.text
        return resultat
    except:
        try:
            resultat =Res.text
        except:
            gpt = " Error !!!"
            resultat = Res.text
        return resultat

